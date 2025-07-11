#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
现代文件转换器 - 核心转换模块

Copyright 2024 现代文件转换器项目

Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at

    http://www.apache.org/licenses/LICENSE-2.0

Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

import os
import shutil
from typing import Optional
from PIL import Image
import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import PyPDF2
import tempfile
import re


class FileConverter:
    def __init__(self):
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'document': ['.pdf', '.docx'],
            'spreadsheet': ['.csv', '.xlsx', '.xls'],
            'markdown': ['.md']
        }
        
    def convert(self, source_path: str, output_path: str, target_format: str) -> bool:
        """
        主转换方法
        """
        try:
            if not os.path.exists(source_path):
                raise FileNotFoundError(f"源文件不存在: {source_path}")
                
            source_ext = os.path.splitext(source_path)[1].lower()
            target_ext = f".{target_format.lower()}"
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                
            # 根据文件类型调用相应的转换方法
            if source_ext in self.supported_formats['image']:
                return self._convert_image(source_path, output_path, target_format)
            elif source_ext in self.supported_formats['document']:
                return self._convert_document(source_path, output_path, target_format)
            elif source_ext in self.supported_formats['spreadsheet']:
                return self._convert_spreadsheet(source_path, output_path, target_format)
            elif target_format.upper() == 'MD' and source_ext == '.pdf':
                return self._pdf_to_markdown(source_path, output_path)
            else:
                raise ValueError(f"不支持的源文件格式: {source_ext}")
                
        except Exception as e:
            print(f"转换错误: {e}")
            return False
            
    def _convert_image(self, source_path: str, output_path: str, target_format: str) -> bool:
        """
        图像格式转换
        """
        try:
            if target_format.upper() == 'PDF':
                return self._image_to_pdf(source_path, output_path)
            else:
                with Image.open(source_path) as img:
                    # 处理RGBA图像转换为RGB
                    if img.mode in ('RGBA', 'LA') and target_format.upper() in ['JPG', 'JPEG']:
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[-1])
                        img = background
                    
                    # 保存为目标格式
                    format_name = target_format.upper()
                    if format_name == 'JPG':
                        format_name = 'JPEG'
                    img.save(output_path, format=format_name)
                    return True
                    
        except Exception as e:
            print(f"图像转换错误: {e}")
            return False
            
    def _convert_document(self, source_path: str, output_path: str, target_format: str) -> bool:
        """
        文档格式转换
        """
        try:
            source_ext = os.path.splitext(source_path)[1].lower()
            
            if source_ext == '.pdf' and target_format.upper() == 'DOCX':
                return self._pdf_to_docx(source_path, output_path)
            elif source_ext == '.docx' and target_format.upper() == 'PDF':
                return self._docx_to_pdf(source_path, output_path)
            else:
                # 如果是相同格式，直接复制
                shutil.copy2(source_path, output_path)
                return True
                
        except Exception as e:
            print(f"文档转换错误: {e}")
            return False
            
    def _convert_spreadsheet(self, source_path: str, output_path: str, target_format: str) -> bool:
        """
        表格格式转换
        """
        try:
            source_ext = os.path.splitext(source_path)[1].lower()
            
            if source_ext == '.csv' and target_format.upper() == 'XLSX':
                df = pd.read_csv(source_path, encoding='utf-8')
                df.to_excel(output_path, index=False)
                return True
            elif source_ext in ['.xlsx', '.xls'] and target_format.upper() == 'CSV':
                df = pd.read_excel(source_path)
                df.to_csv(output_path, index=False, encoding='utf-8')
                return True
            elif source_ext in ['.csv', '.xlsx', '.xls'] and target_format.upper() == 'PDF':
                return self._spreadsheet_to_pdf(source_path, output_path)
            else:
                # 如果是相同格式，直接复制
                shutil.copy2(source_path, output_path)
                return True
                
        except Exception as e:
            print(f"表格转换错误: {e}")
            return False
            
    def _image_to_pdf(self, source_path: str, output_path: str) -> bool:
        """
        图像转PDF
        """
        try:
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            
            with Image.open(source_path) as img:
                # 计算适合页面的尺寸
                page_width, page_height = letter
                img_width, img_height = img.size
                
                # 计算缩放比例
                scale_w = (page_width - 2 * inch) / img_width
                scale_h = (page_height - 2 * inch) / img_height
                scale = min(scale_w, scale_h)
                
                new_width = img_width * scale
                new_height = img_height * scale
                
                # 临时保存调整后的图像
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    resized_img = img.resize((int(new_width), int(new_height)), Image.Resampling.LANCZOS)
                    resized_img.save(tmp.name, 'PNG')
                    
                    # 添加到PDF
                    rl_img = RLImage(tmp.name, width=new_width, height=new_height)
                    story.append(rl_img)
                    
                doc.build(story)
                
                # 清理临时文件
                os.unlink(tmp.name)
                
            return True
            
        except Exception as e:
            print(f"图像转PDF错误: {e}")
            return False
            
    def _pdf_to_docx(self, source_path: str, output_path: str) -> bool:
        """
        PDF转Word（简单文本提取）
        """
        try:
            doc = Document()
            
            with open(source_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():
                        doc.add_paragraph(text)
                        
                    # 添加分页符（除了最后一页）
                    if page_num < len(pdf_reader.pages) - 1:
                        doc.add_page_break()
                        
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"PDF转Word错误: {e}")
            return False
            
    def _docx_to_pdf(self, source_path: str, output_path: str) -> bool:
        """
        Word转PDF
        """
        try:
            doc = Document(source_path)
            
            # 创建PDF文档
            pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # 添加段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    p = Paragraph(paragraph.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
                    
            # 添加表格
            for table in doc.tables:
                data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    data.append(row_data)
                
                if data:
                    t = Table(data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 12))
                    
            pdf_doc.build(story)
            return True
            
        except Exception as e:
            print(f"Word转PDF错误: {e}")
            return False
            
    def _spreadsheet_to_pdf(self, source_path: str, output_path: str) -> bool:
        """
        表格转PDF
        """
        try:
            # 读取数据
            source_ext = os.path.splitext(source_path)[1].lower()
            if source_ext == '.csv':
                df = pd.read_csv(source_path, encoding='utf-8')
            else:
                df = pd.read_excel(source_path)
                
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # 添加标题
            title = Paragraph("数据表格", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # 准备表格数据
            data = [df.columns.tolist()]  # 添加列头
            for index, row in df.iterrows():
                data.append([str(cell) for cell in row.tolist()])
                
            # 创建表格
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8)
            ]))
            
            story.append(t)
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"表格转PDF错误: {e}")
            return False

    def _pdf_to_markdown(self, source_path: str, output_path: str) -> bool:
        """
        PDF转Markdown
        """
        try:
            markdown_content = []
            
            with open(source_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():
                        # 处理文本，转换为Markdown格式
                        processed_text = self._process_text_to_markdown(text)
                        markdown_content.append(processed_text)
                        
                        # 添加分页符（除了最后一页）
                        if page_num < len(pdf_reader.pages) - 1:
                            markdown_content.append("\n\n---\n\n")
                            
            # 写入Markdown文件，使用不同平台的编码方式
            import platform
            if platform.system() == 'Windows':
                # Windows使用UTF-8 with BOM
                with open(output_path, 'w', encoding='utf-8-sig') as md_file:
                    md_file.write('\n'.join(markdown_content))
            else:
                # macOS和Linux使用标准UTF-8
                with open(output_path, 'w', encoding='utf-8') as md_file:
                    md_file.write('\n'.join(markdown_content))
                
            return True
            
        except Exception as e:
            print(f"PDF转Markdown错误: {e}")
            return False
            
    def _process_text_to_markdown(self, text: str) -> str:
        """
        处理文本，转换为Markdown格式
        """
        if not text.strip():
            return ""
            
        # 按行分割
        lines = text.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检测可能的标题
            if self._is_likely_title(line):
                # 判断标题级别
                if len(line) < 30:
                    line = f"# {line}"
                elif len(line) < 50:
                    line = f"## {line}"
                else:
                    line = f"### {line}"
            
            # 检测列表项
            elif self._is_likely_list_item(line):
                if not line.startswith('- '):
                    line = f"- {line}"
            
            processed_lines.append(line)
            
        return '\n\n'.join(processed_lines) if processed_lines else ""
    
    def _is_likely_title(self, line: str) -> bool:
        """
        判断是否可能是标题
        """
        if len(line) > 100:
            return False
            
        # 全大写的短行
        if line.isupper() and len(line) < 80:
            return True
            
        # 首字母大写且单词数量少
        if line[0].isupper() and line.count(' ') < 8:
            # 排除包含数字较多的行（可能是数据）
            digit_count = sum(1 for c in line if c.isdigit())
            if digit_count / len(line) < 0.3:
                return True
                
        return False
    
    def _is_likely_list_item(self, line: str) -> bool:
        """
        判断是否可能是列表项
        """
        # 以数字开头的列表项：1. 或 1) 或 1、
        if re.match(r'^\d+[\.\)\、]\s+', line):
            return True
            
        # 以字母开头的列表项：a. 或 a) 或 A.
        if re.match(r'^[a-zA-Z][\.\)]\s+', line):
            return True
            
        # 以符号开头的列表项：• - * ▪ ◦
        if re.match(r'^[•\-\*▪◦]\s+', line):
            return True
            
        # 以短横线开头
        if line.startswith('- '):
            return True
            
        return False
    
    def get_supported_target_formats(self, source_path: str) -> list:
        """
        根据源文件格式返回支持的目标格式
        """
        if not source_path or not os.path.exists(source_path):
            return []
            
        source_ext = os.path.splitext(source_path)[1].lower()
        
        # 定义转换规则
        conversion_rules = {
            # PDF文件
            '.pdf': ['DOCX', 'MD'],
            
            # Word文档
            '.docx': ['PDF'],
            
            # 图像文件
            '.jpg': ['PNG', 'GIF', 'BMP', 'PDF'],
            '.jpeg': ['PNG', 'GIF', 'BMP', 'PDF'],
            '.png': ['JPG', 'GIF', 'BMP', 'PDF'],
            '.gif': ['JPG', 'PNG', 'BMP', 'PDF'],
            '.bmp': ['JPG', 'PNG', 'GIF', 'PDF'],
            '.tiff': ['JPG', 'PNG', 'GIF', 'BMP', 'PDF'],
            '.webp': ['JPG', 'PNG', 'GIF', 'BMP', 'PDF'],
            
            # 表格文件
            '.csv': ['XLSX', 'PDF'],
            '.xlsx': ['CSV', 'PDF'],
            '.xls': ['CSV', 'XLSX', 'PDF']
        }
        
        return conversion_rules.get(source_ext, [])